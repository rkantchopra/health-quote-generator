import os
import re
import io
from datetime import datetime
from typing import Dict, List, Optional

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# ---------- Styling helpers ----------
def set_cell_bg(cell, color="00A36C"):
    cell._tc.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>'))

def set_white_text(cell):
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.bold = True
            r.font.size = Pt(10.5)

def set_table_borders(table):
    tbl = table._tbl
    tbl_pr = tbl.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        tbl.insert(0, tbl_pr)
    old = tbl_pr.find(qn('w:tblBorders'))
    if old is not None:
        tbl_pr.remove(old)
    borders = OxmlElement('w:tblBorders')
    def make(tag):
        e = OxmlElement(f'w:{tag}')
        e.set(qn('w:val'), 'single')
        e.set(qn('w:sz'), '8')
        e.set(qn('w:color'), '000000')
        return e
    for tag in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        borders.append(make(tag))
    tbl_pr.append(borders)

# ---------- MASTER DATA ----------
MASTER: Dict[str, Dict[str, str]] = {
    "ICICI Lombard – Elevate": {
        "Restoration Benefit":"Unlimited (including for same illness)",
        "NCB Benefit":"Unlimited 100% yearly (No Cap)",
        "Room Rent":"Single Private AC",
        "Pre-Hospitalization":"90 Days","Post-Hospitalization":"180 Days",
        "Day Care Treatments":"All covered","Non-Consumables":"All covered",
        "Hospitalization @ Home":"Up to Sum Assured","Ambulance":"Up to Sum Assured",
        "Air Ambulance":"Up to Sum Assured","AYUSH":"Up to Sum Assured","Organ Donor":"Up to Sum Assured",
        "Modern Treatments":"Up to Sum Assured","2-Hour Hospitalization":"Covered","E-Consultation":"Unlimited",
        "Preventive Health Check-up":"All covered (But Optional)",
        "Maternity":"Optional Rider – 10% Of SA & Max 1 Lakh allowed; waiting period 2 yrs (reducible to 1 yr with rider). Newborn Day 1 (10% SA).",
        "New Born Cover":"Day 1 (10% SA)","Worldwide Cover":"Not Available","Lock-the-Clock Premium Freeze":"Not Applicable",
        "OPD Cover":"Optional Rider – can be added","Priority Claim Desk":"Not Available","Cash+ Wallet":"—",
        "Unique Features":"Unlimited NCB (100% SA increase yearly, no cap), 2-hr hospitalization, child cover till 30 yrs, newborn Day-1 (10% SA)"
    },
    "Niva Bupa – ReAssure 3.0": {
        "Restoration Benefit":"Unlimited Sum Reinstatement — Same illness covered multiple times; cover never ends.",
        "NCB Benefit":"Not Applicable – This is an Unlimited Cover Plan.","Room Rent":"Any Room including Suite – No Limit.",
        "Pre-Hospitalization":"60 Days","Post-Hospitalization":"180 Days","Day Care Treatments":"All Day-Care Procedures covered (no limit).",
        "Non-Consumables":"Yes – All Covered.","Hospitalization @ Home":"Covered up to Sum Insured — if medically advised.",
        "Ambulance":"Covered up to Sum Insured — no per-event cap.","Air Ambulance":"Covered up to ₹5 Lakh.",
        "AYUSH":"Covered up to Sum Assured.","Organ Donor":"Covered up to Sum Assured.","Modern Treatments":"Covered up to Sum Assured.",
        "2-Hour Hospitalization":"Covered — short-stay admission eligible.","E-Consultation":"Unlimited online doctor consultations.",
        "Preventive Health Check-up":"Annual health check-up available.","Maternity":"Not Available","New Born Cover":"Not Available",
        "Worldwide Cover":"Worldwide cover with rider; no India-diagnosis rule when opted.",
        "Lock-the-Clock Premium Freeze":"Premium will not increase until a claim occurs — premium stays locked.",
        "OPD Cover":"Optional Rider – ₹1 Lakh annual OPD (Dental, tests, visits, medicines, gym & physio sessions).",
        "Priority Claim Desk":"Optional Rider – Priority claim handling (HNI/Prime).",
        "Cash+ Wallet":"Cashback reward every claim-free year — usable for renewal, co-pay, OPD.",
        "Unique Features":"Unlimited Sum Insured, Worldwide Cover (no India rule), Prime Member service, No-Claim Discounts added to wallet."
    },
    "Niva Bupa – Aspire Platinum": {
        "Restoration Benefit":"Unlimited (including for same illness)","NCB Benefit":"Up to 5X (300%)","Room Rent":"Single Private AC",
        "Pre-Hospitalization":"60 Days","Post-Hospitalization":"90 Days","Day Care Treatments":"All covered",
        "Non-Consumables":"All covered","Hospitalization @ Home":"Up to Sum Assured","Ambulance":"Up to Sum Assured","Air Ambulance":"Up to Sum Assured",
        "AYUSH":"Up to Sum Assured","Organ Donor":"Up to Sum Assured","Modern Treatments":"Up to Sum Assured",
        "2-Hour Hospitalization":"—","E-Consultation":"Unlimited","Preventive Health Check-up":"Standard Available",
        "Maternity":"Standard ₹12k yearly","New Born Cover":"Available","Worldwide Cover":"Not Available",
        "Lock-the-Clock Premium Freeze":"Premium will not increase until a claim occurs — your premium stays locked.",
        "OPD Cover":"Optional Rider – can be added","Priority Claim Desk":"—","Cash+ Wallet":"—",
        "Unique Features":"Premium lock, child cover up to 60 yrs, NCB up to 5X"
    },
    "Tata AIG – Medicare Select": {
        "Restoration Benefit":"Unlimited (including for same illness)","NCB Benefit":"Sum Assured will increase 100% every year, up to 500% (Super NCB).",
        "Room Rent":"Single Private AC","Pre-Hospitalization":"60 Days","Post-Hospitalization":"90 Days","Day Care Treatments":"All covered",
        "Non-Consumables":"All covered","Hospitalization @ Home":"Up to Sum Assured","Ambulance":"Up to Sum Assured","Air Ambulance":"Up to Sum Assured",
        "AYUSH":"Up to Sum Assured","Organ Donor":"Up to Sum Assured","Modern Treatments":"Up to Sum Assured","2-Hour Hospitalization":"—",
        "E-Consultation":"Unlimited","Preventive Health Check-up":"All covered",
        "Maternity":"Optional Rider – 10% of SA and Max Up to 1 Lakh; waiting period 2 yrs (reducible to 1 yr with rider).",
        "New Born Cover":"Available with rider","Worldwide Cover":"Not Available","Lock-the-Clock Premium Freeze":"Not Applicable",
        "OPD Cover":"Optional Rider – can be added","Priority Claim Desk":"—","Cash+ Wallet":"—",
        "Unique Features":"Salary-linked discounts for salaried persons (7.5%), Super NCB up to 500%"
    },
    "HDFC ERGO – Optima Secure": {
        "Restoration Benefit":"Unlimited (including for same illness)","NCB Benefit":"2X Day 1 (e.g., 10 Lakh SA becomes 20 Lakh from day 1)",
        "Room Rent":"Single Private AC","Pre-Hospitalization":"60 Days","Post-Hospitalization":"180 Days","Day Care Treatments":"All covered",
        "Non-Consumables":"All covered","Hospitalization @ Home":"Up to Sum Assured","Ambulance":"Up to Sum Assured","Air Ambulance":"Up to Sum Assured",
        "AYUSH":"Up to Sum Assured","Organ Donor":"Up to Sum Assured","Modern Treatments":"Up to Sum Assured",
        "2-Hour Hospitalization":"—","E-Consultation":"Unlimited","Preventive Health Check-up":"Standard Available",
        "Maternity":"Not Available","New Born Cover":"Not Available","Worldwide Cover":"Not Available",
        "Lock-the-Clock Premium Freeze":"Not Applicable","OPD Cover":"Optional Rider – can be added","Priority Claim Desk":"—",
        "Cash+ Wallet":"—","Unique Features":"2X Cover from Day 1, deductible on 1st claim, health check-up included"
    },
    "Care Health – Supreme": {
        "Restoration Benefit":"Unlimited (including for same illness)","NCB Benefit":"yearly 100% SA increase up to 600% (6X)","Room Rent":"Single Private AC",
        "Pre-Hospitalization":"60 Days","Post-Hospitalization":"180 Days","Day Care Treatments":"All covered","Non-Consumables":"All covered",
        "Hospitalization @ Home":"Up to Sum Assured","Ambulance":"Up to Sum Assured","Air Ambulance":"Up to Sum Assured",
        "AYUSH":"Up to Sum Assured","Organ Donor":"Up to Sum Assured","Modern Treatments":"Up to Sum Assured",
        "2-Hour Hospitalization":"—","E-Consultation":"Unlimited","Preventive Health Check-up":"All covered",
        "Maternity":"Not Available","New Born Cover":"Not Available","Worldwide Cover":"Not Available",
        "Lock-the-Clock Premium Freeze":"Not Applicable","OPD Cover":"Optional Rider – can be added","Priority Claim Desk":"—",
        "Cash+ Wallet":"—","Unique Features":"NCB up to 600% (6X), all non-consumables covered, health check-up rider"
    }
}

FEATURES: List[str] = [
 "🔄 Restoration Benefit","💰 NCB Benefit","🏠 Room Rent","🏥 Pre-Hospitalization","🩹 Post-Hospitalization",
 "🌞 Day Care Treatments","🛠️ Non-Consumables","🏡 Hospitalization @ Home","🚑 Ambulance","✈️ Air Ambulance",
 "🌿 AYUSH","❤️ Organ Donor","🔬 Modern Treatments","⏰ 2-Hour Hospitalization","📱 E-Consultation",
 "✅ Preventive Health Check-up","🤰 Maternity","👶 New Born Cover","🌍 Worldwide Cover",
 "🔒 Lock-the-Clock Premium Freeze","💊 OPD Cover","🤝 Priority Claim Desk","💳 Cash+ Wallet","✨ Unique Features"
]

LOGO_MAP = {
    "ICICI Lombard – Elevate":"icici_lombard",
    "Niva Bupa – ReAssure 3.0":"niva_reassure3",
    "Niva Bupa – Aspire Platinum":"niva_aspire",
    "Tata AIG – Medicare Select":"tata_aig",
    "HDFC ERGO – Optima Secure":"hdfc_ergo",
    "Care Health – Supreme":"care_health"
}

def map_master(name_text: Optional[str]) -> Optional[str]:
    if not name_text:
        return None
    n = str(name_text).lower().strip()
    if ("black" in n and "variant" in n) or "reassure" in n or "v3.0" in n or "v 3.0" in n or ("3.0" in n and "niva" in n):
        return "Niva Bupa – ReAssure 3.0"
    if "aspire" in n:
        return "Niva Bupa – Aspire Platinum"
    if "icici" in n or "lombard" in n or "elevate" in n:
        return "ICICI Lombard – Elevate"
    if "tata" in n or "aig" in n or "medicare" in n:
        return "Tata AIG – Medicare Select"
    if "hdfc" in n or "ergo" in n or "optima" in n:
        return "HDFC ERGO – Optima Secure"
    if "care" in n or "supreme" in n:
        return "Care Health – Supreme"
    for key in MASTER.keys():
        k = key.lower()
        for token in re.split(r'\W+', n):
            if token and token in k:
                return key
    return None

def find_logo_file(master_key: Optional[str], logo_folder: str) -> Optional[str]:
    if not master_key:
        return None
    base = LOGO_MAP.get(master_key)
    if not base:
        base = "".join(ch if ch.isalnum() else "_" for ch in master_key).lower()
    for ext in (".png", ".jpg", ".jpeg", ".webp"):
        p = os.path.join(logo_folder, base + ext)
        if os.path.exists(p):
            return p
    return None

def has_premium(row: pd.Series) -> bool:
    for c in row.index:
        if "prem" in str(c).lower():
            try:
                if float(row[c]) > 0:
                    return True
            except Exception:
                if str(row[c]).strip().upper() not in ("", "0", "NA"):
                    return True
    return False

# ---------- Core generation ----------
def generate_docx(
    excel_input,                 # bytes or str path
    output_path: str,
    logo_folder: str = "logos",
    filename_hint: Optional[str] = None,
) -> str:
    # Choose engine by extension
    def pick_engine(ext: str):
        ext = (ext or "").lower()
        if ext in (".xlsx", ".xlsm"): return "openpyxl"
        if ext == ".xls": return "xlrd"
        return None

    # Read Excel into memory (BytesIO) to avoid file locks
    if isinstance(excel_input, (bytes, bytearray)):
        engine = pick_engine(os.path.splitext(filename_hint or "")[1])
        xls = pd.ExcelFile(io.BytesIO(excel_input), engine=engine)
    else:
        engine = pick_engine(os.path.splitext(str(excel_input))[1])
        xls = pd.ExcelFile(excel_input, engine=engine)
        
    
    required_sheets = ["Client Details", "Premiums"]
    missing = [s for s in required_sheets if s not in xls.sheet_names]
    if missing:
        raise ValueError(f"Missing sheet(s): {', '.join(missing)}. Found sheets: {', '.join(xls.sheet_names)}")

    client_df = pd.read_excel(xls, sheet_name="Client Details")
    premium_df = pd.read_excel(xls, sheet_name="Premiums")

    # Optional but helpful: required columns
    client_required = ["Client Name", "Relation", "DOB", "Age", "City", "Sum Assured"]
    prem_any = ["Plan Name","Plan","Insurance Company","Insurer","Company","Product"]
    prem_prem_cols = ["1 Yr Premium","2 Yr Premium","3 Yr Premium"]

    def _missing(cols, df_cols):
        return [c for c in cols if c not in df_cols]

    miss_client = _missing(client_required, client_df.columns)
    if miss_client:
        raise ValueError(f"'Client Details' is missing columns: {', '.join(miss_client)}")

    if not any(c in premium_df.columns for c in prem_any):
        raise ValueError(f"'Premiums' must have at least one plan-name-like column: {', '.join(prem_any)}")

    # premium columns are optional, but warn if all missing
    if not any(c in premium_df.columns for c in prem_prem_cols):
        # Not fatal; your code already handles missing/0, but this explains it.
        pass

    
    

    required = ["Client Details", "Premiums"]
    for s in required:
        if s not in xls.sheet_names:
            raise RuntimeError(f"Missing sheet: {s}. Found sheets: {xls.sheet_names}")

    client_df = pd.read_excel(xls, sheet_name="Client Details")
    premium_df = pd.read_excel(xls, sheet_name="Premiums")

    premium_df["HasPremium"] = premium_df.apply(has_premium, axis=1)
    valid_premiums = premium_df[premium_df["HasPremium"]].copy().reset_index(drop=True)

    included_master: List[str] = []
    for _, r in valid_premiums.iterrows():
        raw = None
        for c in ["Plan Name", "Plan", "Insurance Company", "Insurer", "Company", "Product"]:
            if c in r and pd.notna(r[c]) and str(r[c]).strip():
                raw = str(r[c]).strip(); break
        if not raw:
            for v in r:
                try:
                    if pd.notna(v) and str(v).strip():
                        raw = str(v).strip(); break
                except Exception: pass
        mapped = map_master(raw)
        if mapped and mapped not in included_master: included_master.append(mapped)
        elif raw and raw not in included_master: included_master.append(raw)

    if not included_master:
        included_master = list(MASTER.keys())

    # ---------- Build DOCX ----------
    doc = Document()
    doc.add_paragraph("🏥 Health Insurance Quote").runs[0].bold = True
    doc.add_paragraph(f"Prepared by your trusted advisor – {datetime.now().strftime('%d-%m-%Y')}")

    # Client Details
    doc.add_paragraph("\n👤 Client Details").runs[0].bold = True
    ct = doc.add_table(rows=1, cols=7)
    for i, h in enumerate(["Member No.","Name","Relation","DOB","Age","City","Sum Assured"]):
        c = ct.rows[0].cells[i]; c.text = h; set_cell_bg(c,"00A36C"); set_white_text(c)
    set_table_borders(ct)

    if client_df.empty:
        ct.add_row()
    else:
        for idx, r in client_df.iterrows():
            row = ct.add_row().cells
            row[0].text = str(idx+1)
            row[1].text = str(r.get("Client Name",""))
            row[2].text = str(r.get("Relation",""))
            row[3].text = str(r.get("DOB",""))
            row[4].text = str(r.get("Age",""))
            row[5].text = str(r.get("City",""))
            row[6].text = str(r.get("Sum Assured",""))

    # Premium Summary
    doc.add_paragraph("\n💰 Premium Summary").runs[0].bold = True
    pt = doc.add_table(rows=1, cols=4)
    for i, h in enumerate(["Insurer / Plan Name","1 Year Premium","2 Year Premium","3 Year Premium"]):
        c = pt.rows[0].cells[i]; c.text = h; set_cell_bg(c,"00A36C"); set_white_text(c)
    set_table_borders(pt)

    for _, r in valid_premiums.iterrows():
        prow = pt.add_row().cells
        raw_label = None
        for c in ["Plan Name","Plan","Insurance Company","Insurer","Company","Product"]:
            if c in r and pd.notna(r[c]) and str(r[c]).strip():
                raw_label = str(r[c]).strip(); break
        master_key = map_master(raw_label)
        logo_path = find_logo_file(master_key, logo_folder) if master_key else None

        cell = prow[0]
        cell.text = ""
        par = cell.paragraphs[0]
        if logo_path:
            try:
                rn = par.add_run()
                rn.add_picture(logo_path, width=Inches(1.0))
                par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception:
                pass
        p2 = cell.add_paragraph()
        run = p2.add_run(raw_label if raw_label else (master_key if master_key else ""))
        run.bold = True; run.font.size = Pt(11)
        p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        for i,col in enumerate(["1 Yr Premium","2 Yr Premium","3 Yr Premium"], start=1):
            val = r.get(col,"")
            prow[i].text = "" if (pd.isna(val) or str(val).strip() in ("","0")) else str(val)

    # Feature Comparison
    doc.add_paragraph("\n🩺 Feature Comparison (Selected Insurers)").runs[0].bold = True
    ncols = 1 + len(included_master)
    ft = doc.add_table(rows=2, cols=ncols)

    first_row = ft.rows[0].cells
    first_row[0].text = ""; set_cell_bg(first_row[0],"00A36C"); set_white_text(first_row[0])
    hdr_row = ft.rows[1].cells
    hdr_row[0].text = "Feature"; set_cell_bg(hdr_row[0],"00A36C"); set_white_text(hdr_row[0])

    for i, name in enumerate(included_master, start=1):
        hdr_row[i].text = name; set_cell_bg(hdr_row[i],"00A36C"); set_white_text(hdr_row[i])
        logo_cell = first_row[i]; logo_cell.text = ""
        logo_path = find_logo_file(name if name in MASTER else map_master(name), logo_folder)
        if logo_path:
            try:
                rrun = logo_cell.paragraphs[0].add_run()
                rrun.add_picture(logo_path, width=Inches(1.0))
                logo_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            except Exception:
                logo_cell.text = ""

    set_table_borders(ft)

    for feat_with_emoji in FEATURES:
        frow = ft.add_row().cells
        frow[0].text = feat_with_emoji
        set_cell_bg(frow[0],"EAF6EA")
        for p in frow[0].paragraphs:
            for r in p.runs: r.font.bold = True; r.font.size = Pt(10.5)

        feat_key = feat_with_emoji.split(" ",1)[1] if " " in feat_with_emoji else feat_with_emoji
        if feat_key.startswith("️"): feat_key = feat_key[1:]
        feat_key = feat_key.strip()

        for j, name in enumerate(included_master, start=1):
            master_key = name if name in MASTER else map_master(name)
            if master_key and master_key in MASTER:
                frow[j].text = MASTER[master_key].get(feat_key,"")
            else:
                frow[j].text = "<Mapping required>"

        if feat_key == "Unique Features":
            for c in frow: set_cell_bg(c,"FFF68F")

    ft.autofit = False
    try:
        ft.columns[0].width = Inches(2.2)
        for i in range(1, ncols): ft.columns[i].width = Inches(1.4)
    except Exception: pass
    set_table_borders(ft)

    # Advisory
    doc.add_paragraph("\n💬 Advisor’s Recommendation").runs[0].bold = True
    adv = doc.add_table(rows=1, cols=2)
    for i,h in enumerate(["Plan","Why choose this plan (quick points)"]):
        c = adv.rows[0].cells[i]; c.text = h; set_cell_bg(c,"00A36C"); set_white_text(c)
    set_table_borders(adv)

    quick_highlights = {
        "ICICI Lombard – Elevate":["Unlimited NCB growth (no cap)","Newborn Day-1 cover (10% SA)","2-hour hospitalization covered"],
        "Niva Bupa – ReAssure 3.0":["Truly Unlimited Sum Insured","Worldwide cover (rider)","Lock-the-Clock premium freeze"],
        "Niva Bupa – Aspire Platinum":["Premium lock","Child cover up to 60 yrs","NCB up to 5X"],
        "Tata AIG – Medicare Select":["Salary-linked discount (7.5%)","Super NCB up to 500%","Optional maternity/newborn rider"],
        "HDFC ERGO – Optima Secure":["2X cover from Day 1","Health check-up included"],
        "Care Health – Supreme":["NCB up to 600% (6X)","Non-consumables covered","Health check-up rider"]
    }

    for plan in included_master:
        row = adv.add_row().cells
        left = row[0]; left.text = ""
        lp = left.paragraphs[0]
        logo_path = find_logo_file(plan if plan in MASTER else map_master(plan), logo_folder)
        if logo_path:
            try:
                rrun = lp.add_run()
                rrun.add_picture(logo_path, width=Inches(1.0))
                lp.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            except Exception:
                pass
        p_name = left.add_paragraph()
        rn = p_name.add_run(plan); rn.bold = True; rn.font.size = Pt(12)

        right = row[1]; right.text = ""
        pm = plan if plan in MASTER else map_master(plan)
        if pm in quick_highlights:
            for pt in quick_highlights[pm]:
                p = right.add_paragraph(f"• {pt}")
                p.paragraph_format.space_before = Pt(1); p.paragraph_format.space_after = Pt(1)
                for run in p.runs: run.font.size = Pt(10)
        else:
            right.text = "See feature table above."

    adv.autofit = False
    try:
        adv.columns[0].width = Inches(2.0); adv.columns[1].width = Inches(4.0)
    except Exception: pass
    set_table_borders(adv)

    note = doc.add_paragraph()
    note.add_run("\nAdvisor Note: ").bold = True
    note.add_run("Choose the plan matching your family's long-term protection, maternity and travel needs. "
                 "Discuss OPD and worldwide rider options before purchase.").italic = True

    # file name
    try:
        client_name = client_df.loc[0,"Client Name"] if "Client Name" in client_df.columns and not client_df.empty else "Client"
    except Exception:
        client_name = "Client"
    safe_name = "".join(c for c in str(client_name) if c.isalnum() or c in (" ", "_")).strip().replace(" ", "_")
    out_name = output_path or f"Health_Quote_{safe_name}_WithLogos.docx"

    os.makedirs(os.path.dirname(out_name) or ".", exist_ok=True)
    doc.save(out_name)
    return out_name
